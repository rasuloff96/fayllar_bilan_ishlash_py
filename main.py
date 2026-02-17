import os
import platform
import json
import subprocess

# Fayllar famatini ozgartirish uchun kerakli kutubxonalar.
import pandas as pd
from fpdf import FPDF
from docx import Document

# Papka yaratish uchun os modulidan foydalanildi.

folder_name = "chiquvchi_fayllar"
if not os.path.exists(folder_name):
    os.mkdir(folder_name)
    print(f"{folder_name} nomli papka yaratildi.")
else:
    print(f"{folder_name} nomli papka allaqachon mavjud.")

file_path = os.path.join(folder_name, "malumotlar.txt")

files = os.listdir(folder_name)
print("Papkadagi fayllar:" , files)



# JSON faylini o'qish va ma'lumotlarni RAM bo'yicha saralash

with open('pcs.json', "r") as file:
    all_pcs = json.load(file)
    
sorted_pcs = sorted(all_pcs, key=lambda pcs: pcs["ram"], reverse=True)

filtered_folder = "saralangan_malumotlar"



# Yangi papka yaratish yoki mavjudligini tekshirish

if not os.path.exists(filtered_folder):
    os.mkdir(filtered_folder)
    print(f"{filtered_folder} nomli papka yaratildi.")
else:
    print(f"{filtered_folder} nomli papka allaqachon mavjud.")
    
    

# Saralangan ma'lumotlarni yangi papkaga yozish

file_path = os.path.join(filtered_folder, "saralangan_malumotlar.txt")
with open(file_path, "w", encoding="utf-8") as f:
    f.write("Saralangan ma'lumotlar (RAM bo'yicha):\n")
    f.write("=" * 50 + "\n")  # "=" belgisi bilan chiziq chizish
    
    for index , pc in enumerate(sorted_pcs, start=1):
        f.write(f"{index}. Model: {pc['model']}\n")
        f.write(f"   RAM: {pc['ram']} GB\n")
        f.write(f"   Storage: {pc['storage']} GB\n")
        f.write("-" * 50 + "\n") # "-" belgisi bilan chiziq chizish
        
print(f"Saralangan ma'lumotlar '{file_path}' fayliga yozildi.")


output_dir = "boshqa_formatdagi_fayllar"
if not os.path.exists(output_dir):
    os.mkdir(output_dir)
    print(f"{output_dir} nomli papka yaratildi.")
else:
    print(f"{output_dir} nomli papka allaqachon mavjud.")


# Exel fomatiga otkazish

df = pd.DataFrame(sorted_pcs)
excel_path = os.path.join(output_dir, "pc_rating.xlsx")
df.to_excel(excel_path, index=False)


# WORD formatiga otkazish

doc = Document()
doc.add_heading("Kuchli kompyuterlar ro'yxati", 0)
for pc in sorted_pcs:
    doc.add_paragraph(f"Model: {pc['model']}")
    doc.add_paragraph(f"RAM: {pc['ram']} GB")
    doc.add_paragraph(f"Storage: {pc['storage']} GB")
    doc.add_paragraph("-" * 30)
    doc.save(os.path.join(output_dir, "pc_rating.docx"))